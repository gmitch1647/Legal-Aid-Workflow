"""
Reference Indexer — turns .docx files in backend/reference_cases/
into embedded chunks stored in Supabase's `reference_chunks` table.

Workflow
--------
1. Scan `backend/reference_cases/` for .docx files
2. For each file:
   a. Extract text via python-docx
   b. Infer document_type and statute_tags from filename + content
   c. Split into semantic chunks (paragraph-based with size control)
   d. Hash the file — skip re-indexing if unchanged
3. Batch-embed all chunks via Voyage AI
4. Upsert into `reference_chunks` table

Can be triggered manually (POST /draft/reindex) or automatically at
backend startup if the table is empty.
"""

import hashlib
import logging
import re
from pathlib import Path
from typing import Iterable, List, Optional

logger = logging.getLogger(__name__)

REFERENCE_DIR = Path(__file__).resolve().parent.parent / "reference_cases"

# Chunking config
MIN_CHUNK_WORDS = 40        # merge shorter paragraphs upward
TARGET_CHUNK_WORDS = 200    # ideal chunk size
MAX_CHUNK_WORDS = 350       # hard max — split longer paragraphs

# Keywords used to auto-tag chunks during indexing
STATUTE_KEYWORDS = {
    "FCRA": [
        "fcra", "1681", "fair credit reporting", "consumer report",
        "credit bureau", "cra", "reinvestig", "reinsert", "failure to delete",
        "furnisher", "1681s-2", "1681e(b)", "1681i",
    ],
    "FDCPA": [
        "fdcpa", "1692", "fair debt collection", "debt collector",
        "cease and desist", "validation", "1692c", "1692e", "1692f",
    ],
    "TCPA": [
        "tcpa", "47 u.s.c", "227", "autodialer", "prerecorded",
        "telephone consumer protection", "do not call", "dnc",
    ],
    "GA_FBPA": [
        "fbpa", "10-1-390", "fair business practices", "georgia fbpa",
        "o.c.g.a", "deceptive trade",
    ],
    "WILLFUL": [
        "willful", "punitive", "1681n", "knowingly", "reckless disregard",
    ],
}

DEFENDANT_KEYWORDS = {
    "Equifax": ["equifax"],
    "Experian": ["experian"],
    "TransUnion": ["transunion", "trans union"],
    "Chex Systems": ["chex systems", "chex"],
    "Midland": ["midland credit", "midland funding", "midland"],
    "LVNV": ["lvnv"],
    "Resurgent": ["resurgent"],
    "Portfolio Recovery": ["portfolio recovery"],
    "Truist": ["truist", "bb&t", "suntrust"],
    "ED Financial": ["ed financial"],
}

# Document type inference from filename
DOCUMENT_TYPE_KEYWORDS = {
    "motion": ["motion", "mtd", "msj", "dismiss", "summary judgment"],
    "discovery": ["discovery", "interrogator", "rfp", "rfa", "request for production"],
    "response": ["response", "answer", "reply", "counterclaim"],
    "demand_letter": ["demand letter", "demand"],
    "guide": ["guide", "resource", "reference"],
}


# ---------------------------------------------------------------------------
# File-level helpers
# ---------------------------------------------------------------------------


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def _extract_docx_text(path: Path) -> str:
    """Extract plain text from a .docx file."""
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise RuntimeError("python-docx is required for reference indexing") from e

    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also pull text from tables (captions often live in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    return "\n\n".join(paragraphs)


def _infer_metadata(filename: str, text: str) -> dict:
    """Infer document_type, statute_tags, and defendants from filename + content."""
    low_name = filename.lower()
    low_text = text[:5000].lower()  # sample first 5k chars for keyword search

    # Document type — check filename first, then fall back to text hints
    document_type = "complaint"  # default
    for dtype, keywords in DOCUMENT_TYPE_KEYWORDS.items():
        if any(kw in low_name for kw in keywords):
            document_type = dtype
            break
    else:
        if "complaint" not in low_name and ("motion" in low_text[:500]):
            document_type = "motion"

    # Statute tags
    tags: List[str] = []
    combined = low_name + " " + low_text
    for tag, keywords in STATUTE_KEYWORDS.items():
        if any(kw in combined for kw in keywords):
            tags.append(tag)

    # Defendants
    defendants: List[str] = []
    for name, keywords in DEFENDANT_KEYWORDS.items():
        if any(kw in combined for kw in keywords):
            defendants.append(name)

    return {
        "document_type": document_type,
        "statute_tags": tags,
        "defendants": defendants,
    }


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def _split_paragraphs(text: str) -> List[str]:
    # Split on double-newline first, then single if that leaves huge blocks
    parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not parts:
        parts = [p.strip() for p in text.split("\n") if p.strip()]
    return parts


def _word_count(s: str) -> int:
    return len(s.split())


def _chunk_paragraphs(paragraphs: List[str]) -> List[str]:
    """Combine small paragraphs and split large ones to produce chunks
    within the target word-count range."""
    chunks: List[str] = []
    buffer: List[str] = []
    buffer_words = 0

    def flush_buffer():
        nonlocal buffer, buffer_words
        if buffer:
            chunks.append("\n\n".join(buffer).strip())
            buffer = []
            buffer_words = 0

    for para in paragraphs:
        wc = _word_count(para)

        # Very long paragraph — split it independently
        if wc > MAX_CHUNK_WORDS:
            flush_buffer()
            sentences = re.split(r"(?<=[.!?])\s+", para)
            sub_buffer: List[str] = []
            sub_words = 0
            for s in sentences:
                sw = _word_count(s)
                if sub_words + sw > TARGET_CHUNK_WORDS and sub_buffer:
                    chunks.append(" ".join(sub_buffer).strip())
                    sub_buffer = [s]
                    sub_words = sw
                else:
                    sub_buffer.append(s)
                    sub_words += sw
            if sub_buffer:
                chunks.append(" ".join(sub_buffer).strip())
            continue

        # Normal paragraph — try to accumulate into target range
        if buffer_words + wc > MAX_CHUNK_WORDS:
            flush_buffer()

        buffer.append(para)
        buffer_words += wc

        # Flush if we've hit target size (and the current buffer isn't tiny)
        if buffer_words >= TARGET_CHUNK_WORDS:
            flush_buffer()

    flush_buffer()

    # Merge any remaining tiny chunks upward
    merged: List[str] = []
    for chunk in chunks:
        if merged and _word_count(chunk) < MIN_CHUNK_WORDS:
            merged[-1] = merged[-1] + "\n\n" + chunk
        else:
            merged.append(chunk)

    return merged


# ---------------------------------------------------------------------------
# Main indexing functions
# ---------------------------------------------------------------------------


def iter_reference_files() -> Iterable[Path]:
    """Yield every .docx file under backend/reference_cases/ (recursively)."""
    if not REFERENCE_DIR.exists():
        return
    for path in sorted(REFERENCE_DIR.rglob("*.docx")):
        if path.name.startswith("~$"):
            continue  # skip Word lock files
        yield path


def count_indexed_chunks(supabase) -> int:
    try:
        resp = (
            supabase.table("reference_chunks")
            .select("id", count="exact")
            .limit(1)
            .execute()
        )
        return resp.count or 0
    except Exception as e:
        logger.warning(f"Could not count indexed chunks: {e}")
        return 0


def index_all_reference_cases(force: bool = False) -> dict:
    """Scan the reference_cases folder and index any new or changed files.

    Parameters
    ----------
    force : bool
        If True, wipe the existing table first and re-index everything.

    Returns
    -------
    dict
        { files_scanned, files_indexed, files_skipped, chunks_indexed, errors }
    """
    from utils.supabase_client import get_supabase
    from utils.embeddings import embed_documents, is_configured

    supabase = get_supabase()

    if not is_configured():
        return {
            "status": "skipped",
            "reason": "VOYAGE_API_KEY not set — cannot index. Add the env var in Railway.",
        }

    if force:
        logger.info("Force reindex — wiping reference_chunks table")
        try:
            supabase.table("reference_chunks").delete().neq(
                "id", "00000000-0000-0000-0000-000000000000"
            ).execute()
        except Exception as e:
            logger.warning(f"Could not wipe reference_chunks: {e}")

    # Fetch existing file hashes so we can skip unchanged files
    existing_hashes: dict[str, str] = {}
    try:
        resp = (
            supabase.table("reference_chunks")
            .select("source_file, file_hash")
            .execute()
        )
        for row in resp.data or []:
            existing_hashes[row["source_file"]] = row["file_hash"]
    except Exception as e:
        logger.warning(f"Could not fetch existing hashes: {e}")

    result = {
        "files_scanned": 0,
        "files_indexed": 0,
        "files_skipped": 0,
        "chunks_indexed": 0,
        "errors": [],
    }

    for path in iter_reference_files():
        result["files_scanned"] += 1
        filename = path.name

        try:
            file_hash = _sha256_file(path)
        except Exception as e:
            logger.error(f"Could not hash {filename}: {e}")
            result["errors"].append(f"{filename}: hash failed")
            continue

        # Skip unchanged files (unless force)
        if not force and existing_hashes.get(filename) == file_hash:
            result["files_skipped"] += 1
            logger.info(f"Skipping unchanged file: {filename}")
            continue

        try:
            text = _extract_docx_text(path)
        except Exception as e:
            logger.error(f"Could not extract text from {filename}: {e}")
            result["errors"].append(f"{filename}: extract failed")
            continue

        if not text.strip():
            result["errors"].append(f"{filename}: empty content")
            continue

        metadata = _infer_metadata(filename, text)
        paragraphs = _split_paragraphs(text)
        chunks = _chunk_paragraphs(paragraphs)

        if not chunks:
            result["errors"].append(f"{filename}: no chunks produced")
            continue

        # Embed all chunks
        try:
            vectors = embed_documents(chunks)
        except Exception as e:
            logger.error(f"Embedding failed for {filename}: {e}")
            result["errors"].append(f"{filename}: embedding failed ({e})")
            continue

        if len(vectors) != len(chunks):
            result["errors"].append(
                f"{filename}: vector count mismatch ({len(vectors)} vs {len(chunks)})"
            )
            continue

        # Delete existing chunks for this file (if present)
        try:
            supabase.table("reference_chunks").delete().eq(
                "source_file", filename
            ).execute()
        except Exception as e:
            logger.warning(f"Could not delete old chunks for {filename}: {e}")

        # Insert new chunks
        rows = []
        for i, (chunk_text, vector) in enumerate(zip(chunks, vectors)):
            rows.append(
                {
                    "source_file": filename,
                    "chunk_index": i,
                    "chunk_text": chunk_text,
                    "word_count": _word_count(chunk_text),
                    "document_type": metadata["document_type"],
                    "statute_tags": metadata["statute_tags"],
                    "defendants": metadata["defendants"],
                    "embedding": vector,
                    "metadata": {"filename": filename},
                    "file_hash": file_hash,
                }
            )

        try:
            # Insert in batches to avoid payload limits
            batch_size = 50
            for start in range(0, len(rows), batch_size):
                supabase.table("reference_chunks").insert(
                    rows[start : start + batch_size]
                ).execute()
            result["files_indexed"] += 1
            result["chunks_indexed"] += len(rows)
            logger.info(
                f"Indexed {filename}: {len(rows)} chunks, "
                f"type={metadata['document_type']}, tags={metadata['statute_tags']}"
            )
        except Exception as e:
            logger.error(f"Insert failed for {filename}: {e}")
            result["errors"].append(f"{filename}: insert failed ({e})")

    return result
