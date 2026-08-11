"""
Legal Document Formatter — generates court-ready .docx files.

Canonical specifications:
- District-specific font module: Northern District of Georgia uses Times New Roman 14pt
- 480-twip line spacing (true double, LineRuleType.AUTO)
- Pt(12) space before AND after every paragraph (240 twips each)
- 1-inch margins, US Letter page size
- Caption: two-column table with single black borders
- Section headers: bold + centered + underlined
- Count headers: three centered bold underlined lines
- Numbered paragraphs with tab, justified and never fully bold/underlined
- Court header: centered bold, single copy only
"""

import io
import re
import logging

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.complaint_safeguards import assert_docx_safe

logger = logging.getLogger(__name__)

SECTION_HEADERS = {
    "COMPLAINT", "PARTIES", "JURISDICTION AND VENUE",
    "FACTUAL ALLEGATIONS", "DAMAGES",
    "WILLFULNESS OF DEFENDANT'S CONDUCT", "WILLFULNESS OF DEFENDANTS' CONDUCT",
    "FULL FILE DISCLOSURE VIOLATIONS", "CONSUMER STATEMENT VIOLATIONS",
    "REINSERTION VIOLATIONS", "VALIDATION NOTICE VIOLATIONS",
    "MINI-MIRANDA VIOLATIONS", "PRAYER FOR RELIEF", "JURY DEMAND",
    "NATURE OF ACTION", "WILLFULNESS",
}


def _set_line_spacing_480(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")


def _format_paragraph(p, bold=False, centered=False, justified=False,
                      space_before=Pt(12), space_after=Pt(12),
                      indent_left=None, hanging=None, underline=False):
    _set_line_spacing_480(p)
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after

    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justified:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if indent_left:
        p.paragraph_format.left_indent = indent_left
    if hanging:
        p.paragraph_format.first_line_indent = -hanging

    return p


def _add_run(paragraph, text, bold=False, underline=False, font_size=12):
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.underline = underline
    return run


def _set_cell_borders(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "000000")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def _set_cell_width(cell, width_dxa):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    # Remove existing tcW
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _add_cell_paragraph(cell, text, bold=False, font_size=12):
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    _set_line_spacing_480(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    return p


def generate_complaint_docx(
    complaint_text: str,
    plaintiff_name: str = "",
    defendant_names: list = None,
    court: str = "United States District Court, Northern District of Georgia, Atlanta Division",
    jury_demand: bool = True,
) -> io.BytesIO:
    if defendant_names is None:
        defendant_names = []

    # The drafter returns plaintext only.  Remove any legacy court-header text
    # so the renderer remains the single caption authority.
    complaint_text = re.sub(
        r"(?im)^\s*(?:IN THE UNITED STATES DISTRICT COURT|FOR THE NORTHERN DISTRICT OF GEORGIA|ATLANTA DIVISION)\s*$\n?",
        "",
        complaint_text or "",
    )
    is_ndga = "NORTHERN DISTRICT OF GEORGIA" in (court or "").upper()
    font_size = 14 if is_ndga else 12
    canonical_names = {
        "transunion": "Trans Union, LLC",
        "transunion llc": "Trans Union, LLC",
        "trans union llc": "Trans Union, LLC",
    }
    defendant_names = [canonical_names.get(str(name).strip().lower(), str(name).strip()) for name in defendant_names if str(name).strip()]

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(font_size)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # Court header — centered bold, SINGLE COPY
    for line in ["IN THE UNITED STATES DISTRICT COURT",
                 "FOR THE NORTHERN DISTRICT OF GEORGIA",
                 "ATLANTA DIVISION"]:
        p = doc.add_paragraph()
        _format_paragraph(p, centered=True)
        _add_run(p, line, bold=True, font_size=font_size)

    # Caption table
    if plaintiff_name or defendant_names:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Table width — single declaration
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        for existing in tblPr.findall(qn("w:tblW")):
            tblPr.remove(existing)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "9360")
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        _set_cell_width(left_cell, 4680)
        _set_cell_width(right_cell, 4680)

        for cell in [left_cell, right_cell]:
            _set_cell_borders(cell)

        # Left: parties
        pname = plaintiff_name.upper() if plaintiff_name else "PLAINTIFF"
        _add_cell_paragraph(left_cell, f"{pname},", bold=True, font_size=font_size)
        _add_cell_paragraph(left_cell, "", font_size=font_size)
        _add_cell_paragraph(left_cell, "Plaintiff,", font_size=font_size)
        _add_cell_paragraph(left_cell, "", font_size=font_size)
        _add_cell_paragraph(left_cell, "v.", font_size=font_size)
        _add_cell_paragraph(left_cell, "", font_size=font_size)
        for dname in defendant_names:
            _add_cell_paragraph(left_cell, f"{dname.upper()},", bold=True, font_size=font_size)
        suffix = "Defendants." if len(defendant_names) != 1 else "Defendant."
        _add_cell_paragraph(left_cell, "", font_size=font_size)
        _add_cell_paragraph(left_cell, suffix, font_size=font_size)

        # Right: case info
        _add_cell_paragraph(right_cell, "CIVIL ACTION NO.", font_size=font_size)
        _add_cell_paragraph(right_cell, "", font_size=font_size)
        _add_cell_paragraph(right_cell, "____________________________", font_size=font_size)
        _add_cell_paragraph(right_cell, "", font_size=font_size)
        _add_cell_paragraph(right_cell, "", font_size=font_size)
        if jury_demand:
            _add_cell_paragraph(right_cell, "JURY TRIAL DEMANDED", bold=True, font_size=font_size)

    # Parse complaint text
    lines = complaint_text.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        upper = stripped.upper()

        # Section headers — bold centered underlined
        if upper in SECTION_HEADERS or (upper == stripped and len(stripped) > 3 and
            any(kw in upper for kw in ["JURISDICTION", "FACTUAL", "PRAYER",
                "WILLFULNESS", "REINSERTION", "DISCLOSURE", "CONSUMER STATEMENT",
                "VALIDATION", "MINI-MIRANDA"])):
            p = doc.add_paragraph()
            _format_paragraph(p, centered=True)
            _add_run(p, stripped, bold=True, underline=True, font_size=font_size)
            continue

        # Count headers — bold centered underlined
        if upper.startswith("COUNT ") and len(stripped) < 30:
            p = doc.add_paragraph()
            _format_paragraph(p, centered=True)
            _add_run(p, stripped, bold=True, underline=True, font_size=font_size)
            continue

        if "VIOLATION OF" in upper and len(stripped) < 150:
            p = doc.add_paragraph()
            _format_paragraph(p, centered=True)
            _add_run(p, stripped, bold=True, underline=True, font_size=font_size)
            continue

        # Defendant line in count header — bold centered, NO underline
        if stripped.startswith("(") and stripped.endswith(")") and len(stripped) < 200:
            p = doc.add_paragraph()
            _format_paragraph(p, centered=True)
            _add_run(p, stripped, bold=True, font_size=font_size)
            continue

        # Numbered paragraphs — justified with tab
        numbered_match = re.match(r"^(\d+)\.\s*(.*)", stripped)
        if numbered_match:
            num = numbered_match.group(1)
            text = numbered_match.group(2)
            p = doc.add_paragraph()
            _format_paragraph(p, justified=True)
            _add_run(p, f"{num}.\t{text}", font_size=font_size)
            continue

        # Lettered prayer items (A., B., C.)
        letter_match = re.match(r"^([A-Z])\.\s*(.*)", stripped)
        if letter_match and len(stripped) < 500:
            letter = letter_match.group(1)
            text = letter_match.group(2)
            p = doc.add_paragraph()
            _format_paragraph(p, justified=True, indent_left=Inches(0.5))
            _add_run(p, f"{letter}.\t{text}", font_size=font_size)
            continue

        # WHEREFORE
        if upper.startswith("WHEREFORE"):
            p = doc.add_paragraph()
            _format_paragraph(p, justified=True)
            _add_run(p, stripped, font_size=font_size)
            continue

        # COMES NOW
        if upper.startswith("COMES NOW"):
            p = doc.add_paragraph()
            _format_paragraph(p, justified=True)
            _add_run(p, stripped, font_size=font_size)
            continue

        # Respectfully submitted
        if "respectfully submitted" in stripped.lower():
            p = doc.add_paragraph()
            _format_paragraph(p, justified=True)
            _add_run(p, stripped, font_size=font_size)
            continue

        # Signature block lines
        if stripped.startswith("_") or stripped.startswith("Date:") or \
           stripped.startswith("Attorney for") or stripped.startswith("["):
            p = doc.add_paragraph()
            _format_paragraph(p, justified=True)
            _add_run(p, stripped, font_size=font_size)
            continue

        # Regular body paragraph — justified
        p = doc.add_paragraph()
        _format_paragraph(p, justified=True)
        _add_run(p, stripped, font_size=font_size)

    # Ensure that a usable attorney signature block carries the required bar number.
    if "georgia bar no." not in complaint_text.lower():
        p = doc.add_paragraph()
        _format_paragraph(p, justified=True)
        _add_run(p, "Georgia Bar No. [______]", font_size=font_size)

    # The structure validator checks court-header uniqueness and prevents heading
    # formatting from being applied to a whole numbered body paragraph.
    assert_docx_safe(doc)

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
