"""
Document formatter utility - generates Word documents using python-docx.
Produces complaint DOCX and strategy memo DOCX files.
"""

import os
import re
import json
import logging
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from utils.supabase_client import get_supabase

logger = logging.getLogger(__name__)


def _set_document_defaults(doc):
    """Set default document formatting: margins, page size, font."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)


def _add_paragraph(doc, text, bold=False, centered=False, font_size=None, underline=False):
    """Add a formatted paragraph to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = font_size or Pt(12)
    run.bold = bold
    run.underline = underline
    return p


def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f"</w:tcBorders>"
    )
    tcPr.append(tcBorders)


def _create_caption_table(doc, plaintiff_name, defendant_names):
    """Create the two-column caption table for the complaint."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left cell - parties
    left_cell = table.cell(0, 0)
    left_text = (
        f"{plaintiff_name.upper()},\n"
        f"          Plaintiff,\n\n"
        f"     v.\n\n"
    )
    for i, name in enumerate(defendant_names):
        left_text += f"{name.upper()}"
        if i < len(defendant_names) - 1:
            left_text += ",\n"
    left_text += ",\n          Defendant(s)."

    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run(left_text)
    left_run.font.name = "Times New Roman"
    left_run.font.size = Pt(12)
    left_run.bold = True

    # Right cell - case info
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lines = [
        "CASE NO. _____",
        "",
        "Complaint for a civil case",
        "",
        "Jury Trial: \u2612 Yes \u2610 No",
    ]
    right_run = right_para.add_run("\n".join(lines))
    right_run.font.name = "Times New Roman"
    right_run.font.size = Pt(12)

    # Set borders
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell)

    doc.add_paragraph()  # spacing after table
    return table


def _parse_complaint_sections(complaint_text):
    """Parse raw complaint text into structured sections."""
    sections = []
    current_section = {"title": "", "content": ""}

    section_markers = [
        "INTRODUCTION",
        "JURISDICTION",
        "VENUE",
        "PARTIES",
        "FCRA FINDINGS",
        "FACTUAL ALLEGATIONS",
        "COUNT",
        "PRAYER FOR RELIEF",
        "WHEREFORE",
        "TRIAL BY JURY",
        "JURY DEMAND",
        "RESPECTFULLY SUBMITTED",
    ]

    lines = complaint_text.split("\n")
    for line in lines:
        stripped = line.strip()
        upper = stripped.upper()

        is_section_header = False
        for marker in section_markers:
            if upper.startswith(marker):
                is_section_header = True
                break

        if is_section_header:
            if current_section["title"] or current_section["content"].strip():
                sections.append(current_section)
            current_section = {"title": stripped, "content": ""}
        else:
            current_section["content"] += line + "\n"

    if current_section["title"] or current_section["content"].strip():
        sections.append(current_section)

    return sections


def generate_complaint_docx(case_id, complaint_text, case_data, defendants, version=1):
    """Generate a formatted complaint Word document and upload to Supabase Storage."""
    try:
        doc = Document()
        _set_document_defaults(doc)

        # Header - Court name
        court = case_data.get("court", "UNITED STATES DISTRICT COURT\nNORTHERN DISTRICT OF GEORGIA\nATLANTA DIVISION")
        for court_line in court.split("\n"):
            _add_paragraph(doc, court_line.strip(), bold=True, centered=True)

        # Caption table
        plaintiff_name = case_data.get("plaintiff_name", "PLAINTIFF")
        defendant_names = [d.get("name", "DEFENDANT") for d in defendants] if defendants else ["DEFENDANT"]
        _create_caption_table(doc, plaintiff_name, defendant_names)

        # Parse and format sections
        sections = _parse_complaint_sections(complaint_text)

        for section in sections:
            title = section["title"].strip()
            content = section["content"].strip()

            if not title and not content:
                continue

            upper_title = title.upper()

            # Section headers
            if any(marker in upper_title for marker in ["INTRODUCTION", "JURISDICTION", "VENUE", "PARTIES", "FCRA FINDINGS", "FACTUAL ALLEGATIONS", "PRAYER", "WHEREFORE"]):
                _add_paragraph(doc, title, bold=True, centered=True, underline=True)
            elif "COUNT" in upper_title:
                _add_paragraph(doc, title, bold=True, centered=True)
            elif "TRIAL BY JURY" in upper_title or "JURY DEMAND" in upper_title:
                _add_paragraph(doc, "TRIAL BY JURY IS DEMANDED", bold=True, centered=True)
                continue
            elif "RESPECTFULLY" in upper_title:
                _add_paragraph(doc, title, centered=False)

            # Section content
            if content:
                paragraphs = content.split("\n\n")
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue

                    # Check for numbered paragraphs
                    num_match = re.match(r"^(\d+)\.\s*(.*)", para_text, re.DOTALL)
                    if num_match:
                        p = doc.add_paragraph()
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(12)
                        p.paragraph_format.first_line_indent = Inches(0.5)
                        run = p.add_run(f"{num_match.group(1)}. {num_match.group(2).strip()}")
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                    else:
                        _add_paragraph(doc, para_text)

        # Signature block
        doc.add_paragraph()
        _add_paragraph(doc, "Respectfully submitted,")
        doc.add_paragraph()
        _add_paragraph(doc, "Date: ____________________")
        doc.add_paragraph()
        _add_paragraph(doc, f"____________________")
        _add_paragraph(doc, plaintiff_name)

        address = case_data.get("plaintiff_address", "")
        if address:
            _add_paragraph(doc, address)
        phone = case_data.get("plaintiff_phone", "")
        if phone:
            _add_paragraph(doc, phone)
        email = case_data.get("plaintiff_email", "")
        if email:
            _add_paragraph(doc, email)

        # Save to temp file and upload
        storage_path = f"cases/{case_id}/complaint_v{version}.docx"
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            supabase = get_supabase()
            with open(tmp_path, "rb") as f:
                supabase.storage.from_("documents").upload(
                    storage_path,
                    f,
                    {
                        "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "upsert": "true",
                    },
                )
            logger.info(f"Complaint DOCX uploaded to {storage_path}")
        except Exception as e:
            logger.error(f"Failed to upload complaint DOCX: {e}")
        finally:
            os.unlink(tmp_path)

        return storage_path

    except Exception as e:
        logger.error(f"Error generating complaint DOCX: {e}")
        raise


def generate_strategy_memo(case_id, agent_outputs, version=1):
    """Generate a strategy memo Word document and upload to Supabase Storage."""
    try:
        doc = Document()
        _set_document_defaults(doc)

        # Title
        _add_paragraph(doc, "CONFIDENTIAL — ATTORNEY WORK PRODUCT", bold=True, centered=True)
        _add_paragraph(doc, "STRATEGY MEMORANDUM", bold=True, centered=True, font_size=Pt(14))
        doc.add_paragraph()

        # Parse agent outputs
        intake = agent_outputs.get("intake_analyst", {})
        classification = agent_outputs.get("case_classifier", {})
        research = agent_outputs.get("legal_researcher", {})
        damages = agent_outputs.get("damages_analyst", {})
        qa = agent_outputs.get("qa_reviewer", {})

        # Case Overview
        _add_paragraph(doc, "I. CASE OVERVIEW", bold=True, centered=True, underline=True)
        if isinstance(intake, dict):
            plaintiff = intake.get("plaintiff_name", "N/A")
            county = intake.get("county", "N/A")
            _add_paragraph(doc, f"Plaintiff: {plaintiff}")
            _add_paragraph(doc, f"County of Residence: {county}")
            if intake.get("summary"):
                _add_paragraph(doc, f"Summary: {intake['summary']}")
            if intake.get("key_dates"):
                _add_paragraph(doc, "Key Dates:", bold=True)
                dates = intake["key_dates"]
                if isinstance(dates, dict):
                    for k, v in dates.items():
                        _add_paragraph(doc, f"  - {k}: {v}")
                elif isinstance(dates, list):
                    for d in dates:
                        _add_paragraph(doc, f"  - {d}")
        else:
            _add_paragraph(doc, str(intake) if intake else "No intake data available.")

        # Defendants
        _add_paragraph(doc, "II. DEFENDANTS", bold=True, centered=True, underline=True)
        if isinstance(intake, dict) and intake.get("defendants"):
            for d in intake["defendants"]:
                if isinstance(d, dict):
                    name = d.get("name", "Unknown")
                    role = d.get("role", d.get("entity_type", ""))
                    _add_paragraph(doc, f"- {name} ({role})")
                else:
                    _add_paragraph(doc, f"- {d}")
        elif isinstance(classification, dict) and classification.get("defendants"):
            for d in classification["defendants"]:
                if isinstance(d, dict):
                    _add_paragraph(doc, f"- {d.get('name', d.get('defendant', 'Unknown'))}")
                else:
                    _add_paragraph(doc, f"- {d}")

        # Violations Identified
        _add_paragraph(doc, "III. VIOLATIONS IDENTIFIED", bold=True, centered=True, underline=True)
        if isinstance(classification, dict) and classification.get("violations"):
            for v in classification["violations"]:
                if isinstance(v, dict):
                    statute = v.get("statute", v.get("section", ""))
                    desc = v.get("description", v.get("violation", ""))
                    defendant = v.get("defendant", "")
                    _add_paragraph(doc, f"- {statute}: {desc}")
                    if defendant:
                        _add_paragraph(doc, f"  Defendant(s): {defendant}")
                else:
                    _add_paragraph(doc, f"- {v}")
        else:
            _add_paragraph(doc, json.dumps(classification, indent=2) if classification else "No classification data.")

        # Count Structure
        _add_paragraph(doc, "IV. COUNT STRUCTURE", bold=True, centered=True, underline=True)
        if isinstance(research, dict) and research.get("counts"):
            for c in research["counts"]:
                if isinstance(c, dict):
                    count_num = c.get("count_number", c.get("count", ""))
                    title = c.get("title", c.get("statute", ""))
                    defendants_list = c.get("defendants", "")
                    _add_paragraph(doc, f"Count {count_num}: {title}", bold=True)
                    if defendants_list:
                        _add_paragraph(doc, f"  Against: {defendants_list}")
                    if c.get("elements"):
                        _add_paragraph(doc, f"  Elements: {c['elements']}")
                else:
                    _add_paragraph(doc, f"- {c}")
        else:
            _add_paragraph(doc, json.dumps(research, indent=2) if research else "No research data.")

        # Damages Analysis
        _add_paragraph(doc, "V. DAMAGES ANALYSIS", bold=True, centered=True, underline=True)
        if isinstance(damages, dict):
            for key in ["actual_damages", "statutory_damages", "punitive_damages", "attorney_fees", "total_potential"]:
                if damages.get(key):
                    label = key.replace("_", " ").title()
                    _add_paragraph(doc, f"{label}: {damages[key]}")
            if damages.get("per_count"):
                _add_paragraph(doc, "Per-Count Breakdown:", bold=True)
                for pc in damages["per_count"]:
                    if isinstance(pc, dict):
                        _add_paragraph(doc, f"  - {pc.get('count', '')}: {pc.get('damages', '')}")
                    else:
                        _add_paragraph(doc, f"  - {pc}")
        else:
            _add_paragraph(doc, json.dumps(damages, indent=2) if damages else "No damages data.")

        # Recommended Strategy
        _add_paragraph(doc, "VI. RECOMMENDED STRATEGY", bold=True, centered=True, underline=True)
        if isinstance(classification, dict) and classification.get("court"):
            _add_paragraph(doc, f"Recommended Court: {classification['court']}")
        if isinstance(classification, dict) and classification.get("reasoning"):
            _add_paragraph(doc, f"Court Selection Reasoning: {classification['reasoning']}")
        _add_paragraph(doc, "1. File complaint in recommended court.")
        _add_paragraph(doc, "2. Serve all defendants per Federal Rules of Civil Procedure.")
        _add_paragraph(doc, "3. Prepare for discovery including document requests to all CRAs and furnishers.")
        _add_paragraph(doc, "4. Obtain updated credit reports to document ongoing violations.")

        # Risk Assessment
        _add_paragraph(doc, "VII. RISK ASSESSMENT", bold=True, centered=True, underline=True)
        if isinstance(qa, dict):
            if qa.get("approved"):
                _add_paragraph(doc, "QA Review: PASSED - Complaint meets all quality standards.")
            else:
                _add_paragraph(doc, "QA Review: ISSUES IDENTIFIED")
                if qa.get("issues"):
                    for issue in qa["issues"]:
                        _add_paragraph(doc, f"  - {issue}")
        else:
            _add_paragraph(doc, "QA review data not available.")

        # Footer
        doc.add_paragraph()
        _add_paragraph(doc, f"Memo generated: {datetime.now().strftime('%B %d, %Y')}", centered=True)
        _add_paragraph(doc, "CONFIDENTIAL — ATTORNEY WORK PRODUCT", bold=True, centered=True)

        # Save and upload
        storage_path = f"cases/{case_id}/strategy_memo_v{version}.docx"
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            supabase = get_supabase()
            with open(tmp_path, "rb") as f:
                supabase.storage.from_("documents").upload(
                    storage_path,
                    f,
                    {
                        "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "upsert": "true",
                    },
                )
            logger.info(f"Strategy memo uploaded to {storage_path}")
        except Exception as e:
            logger.error(f"Failed to upload strategy memo: {e}")
        finally:
            os.unlink(tmp_path)

        return storage_path

    except Exception as e:
        logger.error(f"Error generating strategy memo: {e}")
        raise


def format_complaint_documents(case_id, complaint_text, case_data, defendants, agent_outputs, version=1):
    """Generate both complaint DOCX and strategy memo, return URLs."""
    complaint_url = generate_complaint_docx(case_id, complaint_text, case_data, defendants, version)
    memo_url = generate_strategy_memo(case_id, agent_outputs, version)

    # Update complaints table with URLs
    try:
        supabase = get_supabase()
        supabase.table("complaints").update({
            "complaint_docx_url": complaint_url,
            "strategy_memo_url": memo_url,
        }).eq("case_id", case_id).eq("is_current", True).execute()
    except Exception as e:
        logger.error(f"Failed to update complaint URLs: {e}")

    return {"complaint_url": complaint_url, "memo_url": memo_url}
