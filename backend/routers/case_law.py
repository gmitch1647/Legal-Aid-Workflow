"""
Case Law router — manages judicial opinions for RAG retrieval.

Supports:
- Uploading opinions (text or .docx)
- AI-powered summarization and tagging
- Vector indexing for semantic search
- Search/retrieval for agent context injection
"""

import logging
import uuid
from datetime import datetime, timezone
from typing import Optional

from fastapi import APIRouter, File, Form, Header, HTTPException, UploadFile
from pydantic import BaseModel

from utils.supabase_client import get_supabase

logger = logging.getLogger(__name__)

router = APIRouter()


async def _get_current_user(authorization: str) -> dict:
    from routers.cases import get_current_user as _shared
    return await _shared(authorization)


def _require_attorney(profile: dict):
    if profile.get("role") != "attorney":
        raise HTTPException(status_code=403, detail="Attorney access required")


# ---------------------------------------------------------------------------
# GET / — list all case law entries
# ---------------------------------------------------------------------------

@router.get("")
async def list_case_law(
    court: Optional[str] = None,
    statute: Optional[str] = None,
    search: Optional[str] = None,
    authorization: str = Header(default=None),
):
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()
    query = (
        supabase.table("case_law")
        .select("id, case_name, citation, court, year, statutes, holding, summary, tags, indexed, created_at")
        .order("year", desc=True)
    )

    if court:
        query = query.eq("court", court)

    resp = query.limit(200).execute()
    results = resp.data or []

    if statute:
        results = [r for r in results if statute in (r.get("statutes") or [])]
    if search:
        term = search.lower()
        results = [
            r for r in results
            if term in (r.get("case_name") or "").lower()
            or term in (r.get("holding") or "").lower()
            or term in (r.get("summary") or "").lower()
            or any(term in t.lower() for t in (r.get("tags") or []))
        ]

    return results


# ---------------------------------------------------------------------------
# GET /{id} — get full case law entry
# ---------------------------------------------------------------------------

@router.get("/{case_id}")
async def get_case_law_entry(case_id: str, authorization: str = Header(default=None)):
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()
    resp = supabase.table("case_law").select("*").eq("id", case_id).execute()
    if not resp.data:
        raise HTTPException(status_code=404, detail="Case not found")
    return resp.data[0]


# ---------------------------------------------------------------------------
# POST / — add case law entry (with optional AI summarization)
# ---------------------------------------------------------------------------

class CaseLawCreate(BaseModel):
    case_name: str
    citation: Optional[str] = ""
    court: Optional[str] = ""
    year: Optional[int] = None
    statutes: list = []
    holding: Optional[str] = ""
    full_text: Optional[str] = ""
    summary: Optional[str] = ""
    tags: list = []
    source_file: Optional[str] = ""


@router.post("")
async def create_case_law(payload: CaseLawCreate, authorization: str = Header(default=None)):
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()
    record = payload.model_dump()
    record["id"] = str(uuid.uuid4())
    record["indexed"] = False
    record["created_at"] = datetime.now(timezone.utc).isoformat()

    resp = supabase.table("case_law").insert(record).execute()
    if not resp.data:
        raise HTTPException(status_code=500, detail="Failed to create case law entry")

    # Auto-index if text is present
    if payload.full_text and len(payload.full_text) > 100:
        import asyncio
        asyncio.create_task(_index_case_law(record["id"], payload.full_text))

    return resp.data[0]


# ---------------------------------------------------------------------------
# POST /upload — upload a case law document (.docx, .txt, .pdf text)
# ---------------------------------------------------------------------------

@router.post("/upload")
async def upload_case_law(
    file: UploadFile = File(...),
    case_name: str = Form(""),
    citation: str = Form(""),
    court: str = Form(""),
    year: int = Form(0),
    statutes: str = Form(""),
    authorization: str = Header(default=None),
):
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    # Read file content
    content = await file.read()
    text = ""

    if file.filename.endswith(".txt"):
        text = content.decode("utf-8", errors="ignore")
    elif file.filename.endswith(".docx"):
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read .docx: {e}")
    elif file.filename.endswith(".pdf"):
        try:
            import io
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read .pdf: {e}")
    else:
        text = content.decode("utf-8", errors="ignore")

    if not text or len(text) < 50:
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    # Parse statutes list
    statute_list = [s.strip() for s in statutes.split(",") if s.strip()] if statutes else []

    supabase = get_supabase()
    record_id = str(uuid.uuid4())

    record = {
        "id": record_id,
        "case_name": case_name or file.filename.replace(".docx", "").replace(".pdf", "").replace(".txt", ""),
        "citation": citation,
        "court": court,
        "year": year if year > 0 else None,
        "statutes": statute_list,
        "holding": "",
        "full_text": text,
        "summary": "",
        "tags": [],
        "source_file": file.filename,
        "indexed": False,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }

    resp = supabase.table("case_law").insert(record).execute()
    if not resp.data:
        raise HTTPException(status_code=500, detail="Failed to save case law")

    # Kick off AI summarization + indexing in background
    import asyncio
    asyncio.create_task(_summarize_and_index(record_id, text, case_name or file.filename))

    return {
        "id": record_id,
        "case_name": record["case_name"],
        "status": "processing",
        "message": "File uploaded. AI is summarizing and indexing for search.",
    }


# ---------------------------------------------------------------------------
# POST /bulk-upload — upload multiple files at once for bulk ingestion
# ---------------------------------------------------------------------------

@router.post("/bulk-upload")
async def bulk_upload_case_law(
    files: list[UploadFile] = File(...),
    authorization: str = Header(default=None),
):
    """Upload multiple legal documents for AI categorization and indexing."""
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    documents = []
    for file in files:
        content = await file.read()
        text = ""

        if file.filename.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
        elif file.filename.endswith(".docx"):
            try:
                import io as _io
                from docx import Document
                doc = Document(_io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception:
                text = content.decode("utf-8", errors="ignore")
        elif file.filename.endswith(".pdf"):
            try:
                import io as _io
                from PyPDF2 import PdfReader
                reader = PdfReader(_io.BytesIO(content))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception:
                pass
        else:
            text = content.decode("utf-8", errors="ignore")

        if text and len(text) >= 50:
            documents.append({"filename": file.filename, "text": text})

    if not documents:
        raise HTTPException(status_code=400, detail="No valid documents found in upload")

    # Process in background
    import asyncio
    from utils.knowledge_importer import bulk_ingest_texts
    asyncio.create_task(bulk_ingest_texts(documents))

    return {
        "status": "processing",
        "files_accepted": len(documents),
        "files_rejected": len(files) - len(documents),
        "message": f"Processing {len(documents)} files. They will appear in the Knowledge Base as they're indexed.",
    }


# ---------------------------------------------------------------------------
# POST /ingest-text — ingest raw text (for pasting or API integration)
# ---------------------------------------------------------------------------

@router.post("/ingest-text")
async def ingest_text(
    request_body: dict,
    authorization: str = Header(default=None),
):
    """Ingest raw text as a knowledge base entry."""
    from fastapi import Request

    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    text = request_body.get("text", "")
    filename = request_body.get("filename", "pasted_document")

    if not text or len(text) < 50:
        raise HTTPException(status_code=400, detail="Text too short")

    import asyncio
    from utils.knowledge_importer import ingest_document
    asyncio.create_task(ingest_document(text, filename))

    return {"status": "processing", "message": "Document is being categorized and indexed."}


# ---------------------------------------------------------------------------
# POST /reprocess-all — fix stuck "Processing" entries
# ---------------------------------------------------------------------------

@router.post("/reprocess-all")
async def reprocess_all_stuck(authorization: str = Header(default=None)):
    """Reprocess all entries stuck in 'Processing' (indexed=false)."""
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()
    resp = supabase.table("case_law").select("id, full_text, case_name").eq("indexed", False).execute()

    if not resp.data:
        return {"status": "none_stuck", "count": 0}

    import asyncio
    count = 0
    for entry in resp.data:
        text = entry.get("full_text", "")
        if text and len(text) > 50:
            asyncio.create_task(_summarize_and_index(entry["id"], text, entry.get("case_name", "")))
            count += 1
        else:
            # No text — just mark as indexed so it stops showing "Processing"
            supabase.table("case_law").update({"indexed": True}).eq("id", entry["id"]).execute()
            count += 1

    return {"status": "reprocessing", "count": count}


# ---------------------------------------------------------------------------
# POST /{id}/reindex — re-index a case law entry
# ---------------------------------------------------------------------------

@router.post("/{case_id}/reindex")
async def reindex_case_law(case_id: str, authorization: str = Header(default=None)):
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()
    resp = supabase.table("case_law").select("full_text, case_name").eq("id", case_id).execute()
    if not resp.data:
        raise HTTPException(status_code=404, detail="Not found")

    text = resp.data[0].get("full_text", "")
    if not text:
        raise HTTPException(status_code=400, detail="No text to index")

    import asyncio
    asyncio.create_task(_summarize_and_index(case_id, text, resp.data[0].get("case_name", "")))

    return {"status": "reprocessing"}


# ---------------------------------------------------------------------------
# DELETE /{id}
# ---------------------------------------------------------------------------

@router.delete("/{case_id}")
async def delete_case_law(case_id: str, authorization: str = Header(default=None)):
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()
    supabase.table("case_law_chunks").delete().eq("case_law_id", case_id).execute()
    supabase.table("case_law").delete().eq("id", case_id).execute()
    return {"deleted": True}


# ---------------------------------------------------------------------------
# GET /search — semantic search across case law
# ---------------------------------------------------------------------------

@router.get("/search/semantic")
async def search_case_law(
    q: str,
    top_k: int = 5,
    authorization: str = Header(default=None),
):
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    try:
        from utils.embeddings import embed_text, is_configured
        if not is_configured():
            raise HTTPException(status_code=400, detail="Embeddings not configured (VOYAGE_API_KEY)")

        query_embedding = embed_text(q)

        supabase = get_supabase()
        resp = supabase.rpc("match_case_law_chunks", {
            "query_embedding": query_embedding,
            "match_threshold": 0.5,
            "match_count": top_k,
        }).execute()

        results = []
        for chunk in (resp.data or []):
            # Get parent case info
            case_resp = supabase.table("case_law").select(
                "case_name, citation, court, year, holding"
            ).eq("id", chunk["case_law_id"]).limit(1).execute()

            case_info = case_resp.data[0] if case_resp.data else {}
            results.append({
                "chunk_content": chunk["content"],
                "similarity": chunk.get("similarity", 0),
                "case_name": case_info.get("case_name", ""),
                "citation": case_info.get("citation", ""),
                "court": case_info.get("court", ""),
                "year": case_info.get("year"),
                "holding": case_info.get("holding", ""),
            })

        return results

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Case law search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Background tasks
# ---------------------------------------------------------------------------

async def _summarize_and_index(case_id: str, text: str, case_name: str):
    """AI-powered summarization then vector indexing."""
    try:
        import anthropic
        client = anthropic.Anthropic()

        # Summarize with Haiku
        response = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=2048,
            messages=[{
                "role": "user",
                "content": f"""Analyze this judicial opinion and provide:
1. HOLDING: The key legal holding in 1-2 sentences
2. SUMMARY: A 3-5 sentence summary of the case, facts, and reasoning
3. STATUTES: List of statutes interpreted (e.g., ["FCRA", "1681e(b)", "1681i(a)"])
4. COURT: The court (e.g., "N.D. Ga.", "11th Cir.", "S.Ct.")
5. YEAR: The year of the decision
6. TAGS: 5-8 search tags

Return as JSON: {{"holding": "...", "summary": "...", "statutes": [...], "court": "...", "year": N, "tags": [...]}}

Opinion text (first 15000 chars):
{text[:15000]}"""
            }],
        )

        import json
        raw = response.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw[raw.index("\n") + 1:]
            if raw.endswith("```"):
                raw = raw[:-3].strip()

        try:
            meta = json.loads(raw)
        except Exception:
            meta = {}

        # Update case law record with AI-generated metadata
        supabase = get_supabase()
        updates = {"updated_at": datetime.now(timezone.utc).isoformat()}
        if meta.get("holding"):
            updates["holding"] = meta["holding"]
        if meta.get("summary"):
            updates["summary"] = meta["summary"]
        if meta.get("statutes"):
            updates["statutes"] = meta["statutes"]
        if meta.get("court"):
            updates["court"] = meta["court"]
        if meta.get("year"):
            updates["year"] = meta["year"]
        if meta.get("tags"):
            updates["tags"] = meta["tags"]

        supabase.table("case_law").update(updates).eq("id", case_id).execute()
        logger.info(f"Case law summarized: {case_name}")

    except Exception as e:
        logger.warning(f"Case law summarization failed for {case_id}: {e}")

    # Always mark as indexed after summarization — even if vector embedding fails,
    # the text content is stored and usable for keyword search
    supabase = get_supabase()
    supabase.table("case_law").update({
        "indexed": True,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }).eq("id", case_id).execute()

    # Try vector indexing (optional — depends on VOYAGE_API_KEY being set)
    await _index_case_law(case_id, text)


async def _index_case_law(case_id: str, text: str):
    """Chunk and embed case law text for vector search."""
    supabase = get_supabase()
    try:
        from utils.embeddings import embed_texts, is_configured
        if not is_configured():
            logger.info("Embeddings not configured — text stored but not vector-indexed")
            return

        # Chunk the text (1500 chars with 200 char overlap)
        chunks = []
        chunk_size = 1500
        overlap = 200
        start = 0

        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            if chunk.strip():
                chunks.append(chunk.strip())
            start = end - overlap

        if not chunks:
            return

        # Embed all chunks
        embeddings = embed_texts(chunks)
        if not embeddings or len(embeddings) != len(chunks):
            logger.error(f"Embedding mismatch for case law {case_id}")
            return

        # Delete existing chunks for this case
        supabase.table("case_law_chunks").delete().eq("case_law_id", case_id).execute()

        # Insert new chunks
        records = []
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            records.append({
                "id": str(uuid.uuid4()),
                "case_law_id": case_id,
                "chunk_index": i,
                "content": chunk,
                "embedding": emb,
                "created_at": datetime.now(timezone.utc).isoformat(),
            })

        # Insert in batches
        batch_size = 20
        for i in range(0, len(records), batch_size):
            batch = records[i:i + batch_size]
            supabase.table("case_law_chunks").insert(batch).execute()

        logger.info(f"Case law vector-indexed: {case_id} ({len(chunks)} chunks)")

    except Exception as e:
        logger.warning(f"Case law vector indexing failed for {case_id}: {e} (text still stored)")
