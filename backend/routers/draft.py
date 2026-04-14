"""
Draft router.

Attorney-initiated complaint drafting flow. Uses the existing agent
pipeline but bypasses the client portal intake — attorneys fill in a
form directly and trigger the 7-agent pipeline against arbitrary
plaintiff information.

Endpoints
---------
POST   /draft/start              — kick off a new draft session
GET    /draft/{session_id}/status — pipeline status for polling
GET    /draft/{session_id}/result — final complaint text + docx URLs
POST   /draft/{session_id}/save   — attach draft to an existing case
"""

import asyncio
import json
import logging
import uuid
from datetime import datetime, timezone
from typing import Optional

from fastapi import (
    APIRouter,
    BackgroundTasks,
    File,
    Form,
    Header,
    HTTPException,
    UploadFile,
    status,
)
from pydantic import BaseModel

from utils.supabase_client import get_supabase

logger = logging.getLogger(__name__)

router = APIRouter()


# ---------------------------------------------------------------------------
# Auth helper (delegates to shared profile auto-create logic)
# ---------------------------------------------------------------------------


async def _get_current_user(authorization: str) -> dict:
    from routers.cases import get_current_user as _shared
    return await _shared(authorization)


def _require_attorney(profile: dict) -> None:
    if profile.get("role") != "attorney":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only attorneys can use the drafting tool.",
        )


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------


class DefendantPayload(BaseModel):
    name: str
    entity_type: Optional[str] = None
    principal_address: Optional[str] = None
    ga_registered_agent: Optional[str] = None
    defendant_id: Optional[str] = None  # set if chosen from known defendants


class DraftStartPayload(BaseModel):
    plaintiff_name: str
    plaintiff_county: str
    defendants: list[DefendantPayload]
    court: Optional[str] = None
    statutes: Optional[str] = None
    case_facts: str
    damages_description: str
    jury_demand: bool = True
    georgia_claims: str = "include"  # include | federal_only | agent_decides
    document_urls: list[str] = []
    mode: str = "fast"  # "fast" (2-call, ~15s) or "thorough" (7-agent, ~90s)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _format_case_facts(payload: DraftStartPayload) -> str:
    """Build an enriched case_facts string that includes plaintiff info,
    court/statute preferences, and damages — so the agent pipeline can
    read everything it needs from one field."""
    parts = [
        "=== PLAINTIFF ===",
        f"Name: {payload.plaintiff_name}",
        f"County of Residence: {payload.plaintiff_county}, Georgia",
        "",
        "=== ATTORNEY PREFERENCES ===",
        f"Preferred Court: {payload.court or 'Agent to recommend'}",
        f"Statutes Requested: {payload.statutes or 'Agent to classify from facts'}",
        f"Jury Demand: {'Yes' if payload.jury_demand else 'No'}",
        f"Georgia State Claims: {payload.georgia_claims}",
        "",
        "=== CASE FACTS ===",
        payload.case_facts,
        "",
        "=== DAMAGES DESCRIBED ===",
        payload.damages_description,
    ]
    return "\n".join(parts)


def _resolve_defendant_id(supabase, defendant: DefendantPayload) -> Optional[str]:
    """Given a defendant payload from the form, return a Supabase
    defendant_id — either an existing one (if defendant_id was supplied
    or if a matching row exists) or a newly-created custom row."""
    if defendant.defendant_id:
        return defendant.defendant_id

    # Try to match by name (case-insensitive)
    try:
        match = (
            supabase.table("defendants")
            .select("id")
            .ilike("name", defendant.name.strip())
            .limit(1)
            .execute()
        )
        if match.data:
            return match.data[0]["id"]
    except Exception as e:
        logger.warning(f"Defendant lookup failed: {e}")

    # Insert a custom defendant
    try:
        insert = (
            supabase.table("defendants")
            .insert(
                {
                    "name": defendant.name.strip(),
                    "principal_address": defendant.principal_address or "",
                    "ga_registered_agent": defendant.ga_registered_agent or "",
                    "entity_type": defendant.entity_type or "Furnisher",
                    "is_custom": True,
                }
            )
            .execute()
        )
        if insert.data:
            return insert.data[0]["id"]
    except Exception as e:
        logger.error(f"Failed to create custom defendant: {e}")

    return None


# ---------------------------------------------------------------------------
# POST /start — kick off a draft session
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# GET /list — list all draft sessions for the attorney
# ---------------------------------------------------------------------------


@router.get("/list")
async def list_drafts(authorization: str = Header(default=None)):
    """Return all draft sessions (cases), newest first.

    Returns a list with session_id, plaintiff name (parsed from case_facts),
    status, latest version, and created_at for each draft.
    """
    supabase = get_supabase()

    try:
        # Fetch all cases (drafts + real cases)
        cases_resp = (
            supabase.table("cases")
            .select("id, status, case_facts, created_at, updated_at, client_id")
            .order("created_at", desc=True)
            .limit(50)
            .execute()
        )
        cases = cases_resp.data or []

        # Enrich each case with defendant names and the latest complaint version
        results = []
        for c in cases:
            # Extract plaintiff name from case_facts header
            plaintiff_name = "Untitled Draft"
            facts = c.get("case_facts", "") or ""
            if "=== PLAINTIFF ===" in facts:
                for line in facts.split("\n"):
                    if line.startswith("Name:"):
                        plaintiff_name = line.replace("Name:", "").strip() or plaintiff_name
                        break

            # Get defendants
            defendants = []
            try:
                cd = (
                    supabase.table("case_defendants")
                    .select("defendant_id")
                    .eq("case_id", c["id"])
                    .execute()
                )
                for row in cd.data or []:
                    d = (
                        supabase.table("defendants")
                        .select("name")
                        .eq("id", row["defendant_id"])
                        .limit(1)
                        .execute()
                    )
                    if d.data:
                        defendants.append(d.data[0]["name"])
            except Exception:
                pass

            # Get current complaint version
            version = None
            try:
                comp = (
                    supabase.table("complaints")
                    .select("version")
                    .eq("case_id", c["id"])
                    .eq("is_current", True)
                    .limit(1)
                    .execute()
                )
                if comp.data:
                    version = comp.data[0].get("version")
            except Exception:
                pass

            results.append({
                "session_id": c["id"],
                "plaintiff_name": plaintiff_name,
                "defendants": defendants,
                "status": c.get("status"),
                "version": version,
                "created_at": c.get("created_at"),
                "updated_at": c.get("updated_at"),
            })

        return results
    except Exception as e:
        logger.exception("Failed to list drafts")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/start", status_code=status.HTTP_202_ACCEPTED)
async def start_draft(
    payload: DraftStartPayload,
    authorization: str = Header(...),
):
    """Create a draft case record and launch the agent pipeline."""
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()

    # Create a case row using the attorney as a placeholder client.
    # The real plaintiff info is embedded in case_facts via _format_case_facts.
    case_facts = _format_case_facts(payload)

    now = datetime.now(timezone.utc).isoformat()
    case_insert = (
        supabase.table("cases")
        .insert(
            {
                "client_id": profile["id"],
                "status": "approved_for_processing",
                "court": payload.court or "",
                "jury_demand": payload.jury_demand,
                "case_facts": case_facts,
                "damages_description": payload.damages_description,
                "created_at": now,
                "updated_at": now,
            }
        )
        .execute()
    )
    if not case_insert.data:
        raise HTTPException(status_code=500, detail="Could not create draft session.")
    case_id = case_insert.data[0]["id"]

    # Link defendants (resolve or create them as needed)
    for defendant in payload.defendants:
        if not defendant.name.strip():
            continue
        def_id = _resolve_defendant_id(supabase, defendant)
        if def_id:
            try:
                supabase.table("case_defendants").insert(
                    {"case_id": case_id, "defendant_id": def_id}
                ).execute()
            except Exception as e:
                logger.warning(f"Could not link defendant {def_id} to case {case_id}: {e}")

    # Attach uploaded document URLs as case_documents rows
    for doc_url in payload.document_urls:
        try:
            supabase.table("case_documents").insert(
                {
                    "case_id": case_id,
                    "file_name": doc_url.split("/")[-1],
                    "file_type": doc_url.split(".")[-1] if "." in doc_url else "bin",
                    "storage_path": doc_url,
                    "document_category": "other",
                    "uploaded_by": profile["id"],
                }
            ).execute()
        except Exception as e:
            logger.warning(f"Could not attach document {doc_url}: {e}")

    # Launch the pipeline as an asyncio background task.
    async def _safe_pipeline(cid: str, mode: str, facts: str, dmg: str) -> None:
        try:
            if mode == "fast":
                logger.info(f"Fast draft starting for case {cid}")
                from agents.fast_drafter import run_fast_draft
                await run_fast_draft(cid, facts, dmg)
                logger.info(f"Fast draft finished for case {cid}")
            else:
                logger.info(f"Thorough pipeline starting for case {cid}")
                from agents.orchestrator import run_pipeline
                await run_pipeline(cid)
                logger.info(f"Thorough pipeline finished for case {cid}")
        except Exception as e:
            logger.exception(f"Pipeline CRASHED for case {cid}: {e}")
            try:
                sb = get_supabase()
                sb.table("cases").update({
                    "status": "error",
                    "revision_notes": f"PIPELINE ERROR: {type(e).__name__}: {e}",
                }).eq("id", cid).execute()
            except Exception:
                pass

    asyncio.create_task(_safe_pipeline(case_id, payload.mode, case_facts, payload.damages_description))

    logger.info(f"Draft session started for case {case_id} by {profile.get('email')}")

    return {
        "session_id": case_id,
        "status": "started",
        "message": "Agent pipeline kicked off. Poll /draft/{id}/status for updates.",
    }


# ---------------------------------------------------------------------------
# POST /upload — upload a single document for a draft
# ---------------------------------------------------------------------------


@router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_draft_document(
    file: UploadFile = File(...),
    authorization: str = Header(...),
):
    """Upload one file to Supabase Storage and return its storage path.

    The caller (the DraftComplaint page) uploads files one by one and
    collects the returned storage paths to include in the final /start
    payload as ``document_urls``.
    """
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()

    # Read the file bytes
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty file.")

    # Build a unique storage path under drafts/
    session_folder = str(uuid.uuid4())
    safe_name = file.filename.replace(" ", "_")
    storage_path = f"drafts/{session_folder}/{safe_name}"

    try:
        supabase.storage.from_("documents").upload(
            storage_path,
            content,
            {"content-type": file.content_type or "application/octet-stream"},
        )
    except Exception as e:
        logger.error(f"Upload failed: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {e}")

    return {
        "storage_path": storage_path,
        "file_name": file.filename,
        "size": len(content),
    }


# ---------------------------------------------------------------------------
# GET /{session_id}/status — pipeline progress for polling
# ---------------------------------------------------------------------------


AGENT_DISPLAY = {
    "intake_analyst": ("Intake Analyst", "Reading documents and extracting facts"),
    "case_classifier": ("Case Classifier", "Identifying violations and statutes"),
    "legal_researcher": ("Legal Researcher", "Pulling statutory language and counts"),
    "damages_analyst": ("Damages Analyst", "Calculating damages and pleading language"),
    "complaint_drafter": ("Complaint Drafter", "Writing the full complaint"),
    "qa_reviewer": ("QA Reviewer", "Checking all counts and citations"),
}

AGENT_ORDER = [
    "intake_analyst",
    "case_classifier",
    "legal_researcher",
    "damages_analyst",
    "complaint_drafter",
    "qa_reviewer",
]


@router.get("/{session_id}/status")
async def draft_status(session_id: str, authorization: str = Header(...)):
    """Return the pipeline status for a draft session."""
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()

    # Fetch case row to get overall status
    case_resp = (
        supabase.table("cases").select("*").eq("id", session_id).limit(1).execute()
    )
    if not case_resp.data:
        raise HTTPException(status_code=404, detail="Draft session not found.")
    case = case_resp.data[0]

    # Fetch agent_outputs for this case
    outputs_resp = (
        supabase.table("agent_outputs")
        .select("*")
        .eq("case_id", session_id)
        .execute()
    )
    outputs_by_name: dict[str, dict] = {
        row["agent_name"]: row for row in (outputs_resp.data or [])
    }

    # Build the structured agents list in pipeline order
    agents_payload = []
    complete_count = 0
    for name in AGENT_ORDER:
        display_name, description = AGENT_DISPLAY[name]
        row = outputs_by_name.get(name, {})
        status_value = row.get("status", "pending")
        if status_value == "complete":
            complete_count += 1

        elapsed = None
        started = row.get("started_at")
        completed = row.get("completed_at")
        if started and completed:
            try:
                t_start = datetime.fromisoformat(started.replace("Z", "+00:00"))
                t_end = datetime.fromisoformat(completed.replace("Z", "+00:00"))
                elapsed = round((t_end - t_start).total_seconds(), 1)
            except Exception:
                pass

        # Build log messages from output_data if present
        log_messages = []
        output_data = row.get("output_data") or {}
        if isinstance(output_data, dict):
            if "summary" in output_data:
                log_messages.append(str(output_data["summary"])[:200])
            else:
                # Add a generic one-liner
                keys = list(output_data.keys())[:3]
                if keys:
                    log_messages.append(f"Produced: {', '.join(keys)}")
        if row.get("error_message"):
            log_messages.append(f"Error: {row['error_message']}")

        agents_payload.append(
            {
                "name": name,
                "display_name": display_name,
                "description": description,
                "status": status_value,
                "started_at": started,
                "completed_at": completed,
                "elapsed_seconds": elapsed,
                "log_messages": log_messages,
                "error_message": row.get("error_message"),
            }
        )

    # Overall status mapping
    case_status = case.get("status", "")
    if case_status in ("draft_ready", "attorney_review", "approved", "filed"):
        overall = "complete"
    elif case_status in ("error",):
        overall = "error"
    elif case_status in ("agents_processing", "approved_for_processing"):
        overall = "running"
    else:
        overall = "pending"

    # If status is still approved_for_processing after 15 seconds,
    # something probably went wrong before the orchestrator could start
    if case_status == "approved_for_processing" and complete_count == 0:
        created = case.get("created_at", "")
        if created:
            try:
                created_dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
                age = (datetime.now(timezone.utc) - created_dt).total_seconds()
                if age > 15:
                    overall = "error"
            except Exception:
                pass

    progress_percent = int((complete_count / len(AGENT_ORDER)) * 100)
    if overall == "complete":
        progress_percent = 100

    # Surface any pipeline error stored in the case record
    pipeline_error = None
    revision_notes = case.get("revision_notes") or ""
    if revision_notes.startswith("PIPELINE ERROR:"):
        pipeline_error = revision_notes.replace("PIPELINE ERROR: ", "")

    return {
        "session_id": session_id,
        "overall_status": overall,
        "progress_percent": progress_percent,
        "case_status": case_status,
        "agents": agents_payload,
        "pipeline_error": pipeline_error,
    }


# ---------------------------------------------------------------------------
# GET /{session_id}/result — final complaint + download URLs
# ---------------------------------------------------------------------------


@router.get("/{session_id}/result")
async def draft_result(session_id: str, authorization: str = Header(...)):
    """Return the completed complaint text and document URLs."""
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()

    # Fetch the current complaint row
    complaint_resp = (
        supabase.table("complaints")
        .select("*")
        .eq("case_id", session_id)
        .eq("is_current", True)
        .limit(1)
        .execute()
    )

    if not complaint_resp.data:
        raise HTTPException(
            status_code=404,
            detail="No complaint has been generated yet for this draft session.",
        )

    complaint = complaint_resp.data[0]

    # Pull summary info from agent outputs
    summary: dict = {
        "statutes_identified": [],
        "defendants": [],
        "counts": [],
        "damages_potential": "",
    }

    try:
        outputs = (
            supabase.table("agent_outputs")
            .select("agent_name, output_data")
            .eq("case_id", session_id)
            .eq("status", "complete")
            .execute()
        )
        for row in outputs.data or []:
            data = row.get("output_data") or {}
            if not isinstance(data, dict):
                continue
            if row["agent_name"] == "case_classifier":
                summary["statutes_identified"] = data.get("statutes", []) or data.get(
                    "violations", []
                )
            elif row["agent_name"] == "legal_researcher":
                summary["counts"] = data.get("counts", []) or data.get(
                    "count_structure", []
                )
            elif row["agent_name"] == "damages_analyst":
                summary["damages_potential"] = data.get("total_damages", "") or data.get(
                    "summary", ""
                )

        # Defendants
        cd_resp = (
            supabase.table("case_defendants")
            .select("defendant_id")
            .eq("case_id", session_id)
            .execute()
        )
        for row in cd_resp.data or []:
            d_resp = (
                supabase.table("defendants")
                .select("name")
                .eq("id", row["defendant_id"])
                .limit(1)
                .execute()
            )
            if d_resp.data:
                summary["defendants"].append(d_resp.data[0]["name"])
    except Exception as e:
        logger.warning(f"Could not build draft summary: {e}")

    # Sign URLs for download
    def _sign_url(path: Optional[str]) -> Optional[str]:
        if not path:
            return None
        try:
            # Use signed URL (valid for 1 hour) so the frontend can trigger download
            res = supabase.storage.from_("documents").create_signed_url(path, 3600)
            return res.get("signedURL") or res.get("signed_url") or res.get("signedUrl")
        except Exception as e:
            logger.warning(f"Could not sign url for {path}: {e}")
            return None

    return {
        "session_id": session_id,
        "complaint_text": complaint.get("complaint_text", ""),
        "complaint_docx_url": _sign_url(complaint.get("complaint_docx_url")),
        "strategy_memo_url": _sign_url(complaint.get("strategy_memo_url")),
        "version": complaint.get("version", 1),
        "case_summary": summary,
    }


# ---------------------------------------------------------------------------
# POST /{session_id}/save — attach draft to a client case
# ---------------------------------------------------------------------------


class SaveDraftPayload(BaseModel):
    client_id: Optional[str] = None
    new_case_notes: Optional[str] = None


@router.post("/{session_id}/save")
async def save_draft_to_case(
    session_id: str,
    payload: SaveDraftPayload,
    authorization: str = Header(...),
):
    """Assign the draft case to a real client (moves it into the
    regular case pipeline under that client). If no client_id is
    provided, the draft stays under the attorney as-is."""
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()

    update_fields: dict = {
        "status": "attorney_review",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    if payload.client_id:
        update_fields["client_id"] = payload.client_id

    supabase.table("cases").update(update_fields).eq("id", session_id).execute()

    return {"session_id": session_id, "saved": True}


# ---------------------------------------------------------------------------
# GET /{session_id}/download — generate and download a formatted .docx
# ---------------------------------------------------------------------------


@router.get("/{session_id}/download")
async def download_complaint_docx(
    session_id: str,
    authorization: str = Header(...),
):
    """Generate a properly formatted Word document from the current
    complaint text and return it as a downloadable file.

    Format: Times New Roman 12pt, double line spacing (2.0),
    1-inch margins, US Letter page size.
    """
    import io
    import tempfile

    from fastapi.responses import StreamingResponse

    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    supabase = get_supabase()

    # Fetch current complaint text
    complaint_resp = (
        supabase.table("complaints")
        .select("complaint_text, version")
        .eq("case_id", session_id)
        .eq("is_current", True)
        .limit(1)
        .execute()
    )
    if not complaint_resp.data:
        raise HTTPException(status_code=404, detail="No complaint found for this session.")

    complaint_text = complaint_resp.data[0].get("complaint_text", "")
    version = complaint_resp.data[0].get("version", 1)

    if not complaint_text.strip():
        raise HTTPException(status_code=400, detail="Complaint text is empty.")

    # Generate the Word document
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

        doc = Document()

        # Page setup — US Letter, 1-inch margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        # Default style — Times New Roman 12pt, double spacing
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # Parse complaint text into paragraphs and apply formatting
        lines = complaint_text.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                # Empty line — add a blank paragraph
                doc.add_paragraph("")
                continue

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

            # Detect headers and center/bold them
            upper = stripped.upper()
            is_header = (
                stripped == upper
                and len(stripped) > 3
                and not stripped[0].isdigit()
                and any(kw in upper for kw in [
                    "INTRODUCTION", "JURISDICTION", "PARTIES",
                    "FACTUAL ALLEGATIONS", "PRAYER", "RELIEF",
                    "TRIAL BY JURY", "COUNT", "VIOLATION",
                    "FINDINGS", "SIGNATURE",
                ])
            )

            is_count_header = upper.startswith("COUNT ") or "VIOLATION OF" in upper

            if is_header or is_count_header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            elif stripped == "TRIAL BY JURY IS DEMANDED":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            else:
                run = p.add_run(stripped)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"complaint_v{version}.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
            },
        )

    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed on the server.",
        )
    except Exception as e:
        logger.exception(f"Failed to generate DOCX for session {session_id}")
        raise HTTPException(
            status_code=500,
            detail=f"Document generation failed: {type(e).__name__}: {e}",
        )


# ---------------------------------------------------------------------------
# POST /reindex — rebuild the reference_chunks table from backend/reference_cases/
# ---------------------------------------------------------------------------


class ReindexPayload(BaseModel):
    force: bool = False


@router.post("/reindex")
async def reindex_reference_cases(
    payload: ReindexPayload = ReindexPayload(),
    authorization: str = Header(default=None),
):
    """Rebuild the RAG reference index from .docx files in
    backend/reference_cases/.

    Parameters
    ----------
    force : bool
        If True, wipes the existing index first and re-indexes
        everything. Otherwise only processes files whose hash has
        changed since the last index.

    Returns a summary dict with counts of files processed and any
    errors encountered. Attorney-only.
    """
    profile = await _get_current_user(authorization)
    _require_attorney(profile)

    try:
        from utils.reference_indexer import index_all_reference_cases
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Reference indexer could not be loaded: {e}",
        )

    logger.info(f"Reindex triggered by {profile.get('email')} (force={payload.force})")
    result = index_all_reference_cases(force=payload.force)

    return {
        "status": "done",
        **result,
    }


# ---------------------------------------------------------------------------
# GET /reindex/status — count of indexed chunks + list of indexed files
# ---------------------------------------------------------------------------


@router.get("/reindex/status")
async def reindex_status(authorization: str = Header(default=None)):
    """Return the current state of the RAG reference index."""
    try:
        from utils.reference_indexer import iter_reference_files, count_indexed_chunks, _get_reference_dir
        from utils.embeddings import is_configured
        from pathlib import Path
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Indexer unavailable: {e}")

    supabase = get_supabase()

    # Files on disk
    disk_files = [p.name for p in iter_reference_files()]

    # Files currently indexed (distinct source_file values)
    indexed_files: list[str] = []
    try:
        resp = (
            supabase.table("reference_chunks")
            .select("source_file")
            .execute()
        )
        indexed_files = sorted({row["source_file"] for row in (resp.data or [])})
    except Exception as e:
        logger.warning(f"Could not fetch indexed files: {e}")

    return {
        "voyage_configured": is_configured(),
        "files_on_disk": disk_files,
        "files_indexed": indexed_files,
        "total_chunks": count_indexed_chunks(supabase),
        "missing_from_index": [f for f in disk_files if f not in indexed_files],
        "reference_dir": str(_get_reference_dir()),
        "cwd": str(Path.cwd()),
    }


# ---------------------------------------------------------------------------
# POST /{session_id}/revise — chat-style revision of a drafted complaint
# ---------------------------------------------------------------------------


class RevisePayload(BaseModel):
    message: str
    complaint_text: str  # the current complaint text to revise
    attachment_paths: list[str] = []  # Supabase storage paths to analyze


@router.post("/{session_id}/revise")
async def revise_draft(
    session_id: str,
    payload: RevisePayload,
    authorization: str = Header(default=None),
):
    """Send a revision instruction and get back the updated complaint.

    Uses RAG retrieval + conversation memory + explicit statute guidance
    for high-quality revisions that match the attorney's filed case style.
    """
    import anthropic

    supabase = get_supabase()
    client = anthropic.Anthropic()

    # Load revision history for this session (stored in session_metadata jsonb)
    case_resp = (
        supabase.table("cases")
        .select("revision_history")
        .eq("id", session_id)
        .limit(1)
        .execute()
    )
    conversation_history = []
    if case_resp.data and case_resp.data[0].get("revision_history"):
        conversation_history = case_resp.data[0]["revision_history"] or []

    # Retrieve relevant reference chunks via RAG
    reference_context = ""

    # Read any uploaded attachments
    attachment_text = ""
    if payload.attachment_paths:
        try:
            from utils.document_reader import read_document
            attachment_parts = []
            for path in payload.attachment_paths:
                try:
                    file_type = path.split(".")[-1].lower() if "." in path else "bin"
                    text = read_document(path, file_type)
                    if text:
                        file_name = path.split("/")[-1]
                        # Truncate very long documents
                        if len(text) > 8000:
                            text = text[:8000] + "\n\n[...document truncated, showing first 8000 chars]"
                        attachment_parts.append(f"\n=== ATTACHMENT: {file_name} ===\n{text}\n=== END {file_name} ===\n")
                except Exception as e:
                    logger.warning(f"Could not read attachment {path}: {e}")
                    attachment_parts.append(f"\n[Failed to read {path}: {e}]\n")
            attachment_text = "\n".join(attachment_parts)
            logger.info(f"[revise] Attached {len(payload.attachment_paths)} files ({len(attachment_text)} chars)")
        except Exception as e:
            logger.error(f"[revise] Attachment processing failed: {e}")
    try:
        from utils.rag_retrieval import retrieve_relevant_chunks, format_retrieved_context
        from utils.embeddings import is_configured
        if is_configured():
            query = f"{payload.message}\n\n{payload.complaint_text[:2000]}"
            chunks = retrieve_relevant_chunks(query, top_k=6, document_type="complaint")
            if chunks:
                reference_context = format_retrieved_context(chunks)
                logger.info(f"[revise] RAG retrieved {len(chunks)} chunks")
    except Exception as e:
        logger.warning(f"[revise] RAG retrieval failed: {e}")

    # Build the enhanced system prompt
    system_prompt = """You are a legal complaint revision assistant specializing in consumer protection law in the Northern District of Georgia.

CRITICAL — USE ONLY THESE STATUTES FOR COUNTS:

FOR CRA DEFENDANTS (Equifax, Experian, TransUnion, Chex Systems):
- 15 U.S.C. § 1681e(b) — failure to follow reasonable procedures for maximum accuracy
- 15 U.S.C. § 1681i(a)(1)(A) — failure to conduct reasonable reinvestigation after dispute
- 15 U.S.C. § 1681i(a)(2)(A) — failure to forward all relevant dispute info to furnisher
- 15 U.S.C. § 1681i(a)(4) — failure to review all relevant information submitted by consumer
- 15 U.S.C. § 1681i(a)(5)(A) — failure to promptly delete inaccurate information after reinvestigation
- 15 U.S.C. § 1681i(a)(5)(B)(i)(ii)(iii) and (C) — reinsertion without certification/notice/procedures
- 15 U.S.C. § 1681g — failure to provide full file disclosure
- 15 U.S.C. § 1681i(c) — failure to add consumer statement

FOR FURNISHERS (ED Financial, Truist, debt servicers):
- 15 U.S.C. § 1681s-2(b) — failure to investigate after CRA notice
  (NEVER use §1681s-2(a) — no private right of action)

FOR DEBT COLLECTORS (Midland, LVNV, Portfolio Recovery):
- 15 U.S.C. § 1692c(c), § 1692e, § 1692f, § 1692g

FOR TCPA:
- 47 U.S.C. § 227(b)(1)(A)(iii), § 227(b)(1)(B), § 227(c)

GEORGIA FBPA (only if willful/deceptive conduct):
- O.C.G.A. § 10-1-390 et seq. (NEVER § 34-6-2)

NEVER use: §1681e(d), §1681s-2(a), §1681i(a)(1)(B), O.C.G.A. § 34-6-2

RULES FOR REVISIONS:
1. Return the COMPLETE revised complaint text — not just the changed parts
2. CONDENSE counts: same violation by multiple defendants = ONE count naming all
3. NEVER add parties the attorney didn't specifically request
4. NEVER hallucinate defendants, case numbers, or facts not in the current complaint
5. Count headers: 3 centered lines — "Count [Roman]" / "Violation of the Fair Credit Reporting Act" / "15 U.S.C. § [section] ([Defendants])"
6. Each count must include: realleges paragraph, violation facts, EXACT damages language, EXACT willful/negligent closing
7. Number ALL paragraphs sequentially 1 through end
8. If renumbering is needed after adding/removing a count, renumber ALL affected paragraphs
9. NO markdown — plain text only. No ## headers, no --- dividers, no ** bold markers
10. Plaintiff referenced as "Mr./Ms. [Last Name]" in counts, not "Plaintiff"

STANDARD DAMAGES LANGUAGE (use verbatim in every count):
"As a result of each Defendant's violations of [section], [Plaintiff] suffered actual damages, including but not limited to: loss of credit, denial of credit, loss of ability to purchase or benefit from credit, loss of time due to learning how to defend against the Defendant's violation of his/her rights, damage to reputation from brandishing an inaccurate consumer report to third parties which in turn led to humiliation and embarrassment, anxiety and other mental, physical, and emotional distress."

STANDARD WILLFUL/NEGLIGENT CLOSING (use verbatim in every count):
"The violations by each defendant were willful rendering the Defendant liable for punitive damages in an amount to be determined by the court pursuant to 15 U.S.C § 1681n. In the alternative, each defendant was negligent, which entitles [Plaintiff] to recovery under 15 U.S.C § 1681o. [Plaintiff] is entitled to recover actual damages, statutory damages, cost and attorney's fees from each defendant in an amount to be determined by the court pursuant to 15 U.S.C §§ 1681n and 1681o."

RESPONSE FORMAT:
CHANGES MADE:
- [bullet list of what you changed]

REVISED COMPLAINT:
[full complaint text here]"""

    # Build messages with conversation history
    messages = []

    # Add prior revision turns
    for turn in conversation_history[-6:]:  # last 3 exchanges max
        if turn.get("role") and turn.get("content"):
            messages.append({"role": turn["role"], "content": turn["content"]})

    # Current revision request
    user_content_parts = [
        f"CURRENT COMPLAINT:\n\n"
        f"---BEGIN COMPLAINT---\n{payload.complaint_text}\n---END COMPLAINT---",
    ]
    if attachment_text:
        user_content_parts.append(
            f"\nATTACHMENTS (the attorney uploaded these for you to analyze):\n{attachment_text}"
        )
    user_content_parts.append(
        f"\nAttorney's instruction: {payload.message}\n\n"
        f"Remember: do not add defendants, parties, or facts not requested. "
        f"Return the complete revised complaint."
    )
    user_content = "\n".join(user_content_parts)
    messages.append({"role": "user", "content": user_content})

    # Build system blocks with cached core + optional RAG
    system_blocks = [
        {
            "type": "text",
            "text": system_prompt,
            "cache_control": {"type": "ephemeral"},
        }
    ]
    if reference_context:
        system_blocks.append({"type": "text", "text": reference_context})

    try:
        response = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=16384,
            system=system_blocks,
            messages=messages,
        )

        assistant_text = response.content[0].text

        # Try to extract the revised complaint from the response
        revised_complaint = assistant_text
        if "REVISED COMPLAINT:" in assistant_text:
            parts = assistant_text.split("REVISED COMPLAINT:", 1)
            changes_summary = parts[0].strip()
            revised_complaint = parts[1].strip()
        elif "---BEGIN COMPLAINT---" in assistant_text:
            start = assistant_text.index("---BEGIN COMPLAINT---") + len("---BEGIN COMPLAINT---")
            end = assistant_text.index("---END COMPLAINT---") if "---END COMPLAINT---" in assistant_text else len(assistant_text)
            revised_complaint = assistant_text[start:end].strip()
            changes_summary = assistant_text[:assistant_text.index("---BEGIN COMPLAINT---")].strip()
        else:
            changes_summary = ""

        # Save the revised complaint as a new version
        version_query = (
            supabase.table("complaints")
            .select("version")
            .eq("case_id", session_id)
            .order("version", desc=True)
            .limit(1)
            .execute()
        )
        next_version = 1
        if version_query.data:
            next_version = version_query.data[0]["version"] + 1
            supabase.table("complaints").update(
                {"is_current": False}
            ).eq("case_id", session_id).execute()

        supabase.table("complaints").insert({
            "case_id": session_id,
            "complaint_text": revised_complaint,
            "version": next_version,
            "is_current": True,
        }).execute()

        # Save to revision history for multi-turn memory (try/except — column may not exist yet)
        try:
            updated_history = conversation_history + [
                {"role": "user", "content": payload.message[:2000]},
                {"role": "assistant", "content": (changes_summary or "Revised complaint")[:1000]},
            ]
            # Keep only last 20 turns to avoid unbounded growth
            updated_history = updated_history[-20:]
            supabase.table("cases").update({
                "revision_history": updated_history,
            }).eq("id", session_id).execute()
        except Exception as e:
            logger.warning(f"Could not save revision history: {e}")

        return {
            "revised_complaint": revised_complaint,
            "changes_summary": changes_summary,
            "version": next_version,
            "full_response": assistant_text,
        }

    except Exception as e:
        logger.exception(f"Revision failed for session {session_id}")
        raise HTTPException(
            status_code=500,
            detail=f"Revision failed: {type(e).__name__}: {e}",
        )
