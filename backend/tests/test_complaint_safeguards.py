import io
import unittest

from docx import Document

from utils.complaint_safeguards import (
    ComplaintValidationError,
    assert_complaint_safe,
    audit_credit_report,
    build_drafting_context,
    redact_for_complaint,
    validate_complaint_text,
)
from utils.docx_formatter import generate_complaint_docx


class ComplaintSafeguardTests(unittest.TestCase):
    def test_v01_redacts_ssn_dob_and_account_identifiers_before_drafting(self):
        safe = redact_for_complaint({
            "social_security_number": "123-45-0047",
            "date_of_birth": "11/20/1990",
            "narrative": "Account number 1234567890 and SSN 123-45-0047.",
        })
        self.assertEqual(safe["social_security_number"], "number ending in 0047")
        self.assertEqual(safe["date_of_birth"], "Plaintiff was born in 1990")
        self.assertNotIn("123-45-0047", safe["narrative"])
        self.assertIn("ending in 7890", safe["narrative"])
        with self.assertRaises(ComplaintValidationError) as blocked:
            assert_complaint_safe("1. Plaintiff SSN is 123-45-0047.")
        self.assertTrue(any(issue.startswith("V-01") for issue in blocked.exception.issues))

    def test_v02_blocks_silent_venue_conflict(self):
        context = build_drafting_context(
            {"consumer_state": "FL", "consumer_address": "Miramar, Florida"},
            {"district": "Northern District of Georgia"},
        )
        issues = validate_complaint_text(
            "1. Plaintiff resides in Fulton County, Georgia.\nCOUNT I\n§ 1681e(b) and § 1681n\n",
            context=context,
        )
        self.assertTrue(any(issue.startswith("V-02") for issue in issues))

    def test_v04_blocks_statute_that_appears_only_in_prayer(self):
        issues = validate_complaint_text(
            "1. Factual allegation.\nCOUNT I\nViolation of § 1681e(b).\n"
            "PRAYER FOR RELIEF\nPlaintiff seeks relief under § 1681g(a)(1).\n"
        )
        self.assertTrue(any(issue.startswith("V-04") for issue in issues))

    def test_v05_v06_v07_block_bad_statutory_quote_and_unpaired_counts(self):
        text = '''
1. Factual allegation.
COUNT I
Violation of 15 U.S.C. § 1681i(a)(4)
"The agency shall provide the consumer with written notice of the results."
Plaintiff incorporates paragraphs 1 through 1.
15 U.S.C. § 1681n.
COUNT II
Violation of 15 U.S.C. § 1681e(b)
Plaintiff incorporates paragraphs 1 through 1.
15 U.S.C. § 1681o. Plaintiff seeks statutory damages and punitive damages.
'''
        issues = validate_complaint_text(text)
        self.assertTrue(any(issue.startswith("V-05") for issue in issues))
        self.assertTrue(any(issue.startswith("V-06") for issue in issues))
        self.assertTrue(any(issue.startswith("V-07") for issue in issues))

    def test_brown_trans_union_audit_regression_detects_core_findings(self):
        # Structured representative fixture for the defect audit.  It protects
        # the named findings without retaining the consumer's actual PII.
        tradelines = [
            {
                "name": "Possible Financial JRTK", "balance": "$548", "high_balance": "$450",
                "payment_received": "$0", "last_payment_made": "05/01/2024",
                "pay_status": "was 60 days past due", "date_opened": "01/01/2020",
                "removal_date": "01/01/2025", "adverse_information": True,
                "payment_history": [{"month": "01/2024", "rating": "30"}],
            },
            {
                "name": "CCBank Integra", "payment_received": "$0", "pay_status": "Paid in Full",
                "date_opened": "01/01/2020", "removal_date": "01/01/2025", "adverse_information": True,
                "payment_history": [{"month": "01/2024", "rating": "60"}],
            },
            {
                "name": "PerPay", "payment_received": "$0", "pay_status": "Current Account",
                "adverse_information": True,
                "payment_history": [{"month": "04/2026", "rating": "30"}, {"month": "05/2026", "rating": "60", "payment": "$202"}],
            },
            {
                "name": "Kashable", "payment_received": "$0", "pay_status": "Paid as agreed",
                "adverse_information": True,
                "payment_history": [{"month": "06/2024", "rating": "120"}, {"month": "07/2024", "rating": "OK", "amount_paid": "$0"}],
            },
            {
                "name": "Cap One Auto", "payment_received": "$0", "pay_status": "Paid as agreed",
                "adverse_information": True,
                "payment_history": [{"month": "01/2020", "rating": "RPO"}, {"month": "02/2020", "rating": "OK"}],
            },
            {
                "name": "Sunrise", "payment_received": "$0", "pay_status": "Paid as agreed",
                "adverse_information": True,
                "payment_history": [{"month": "01/2024", "rating": "60"}],
            },
            {
                "name": "Possible Financial BAHH", "date_opened": "01/01/2020", "removal_date": "01/01/2025",
                "adverse_information": True, "payment_history": [{"month": "01/2024", "rating": "30"}],
            },
            {
                "name": "BlinCLoans", "date_opened": "01/01/2020", "removal_date": "01/01/2025",
                "adverse_information": True, "payment_history": [{"month": "01/2024", "rating": "30"}],
            },
        ]
        findings = audit_credit_report({"credit_report": {"tradelines": tradelines}})
        identifiers = [finding["id"] for finding in findings]
        for expected in ("C-01", "C-02", "C-04", "C-06", "C-09", "C-10", "C-17"):
            self.assertIn(expected, identifiers)
        self.assertGreaterEqual(identifiers.count("IMPOSSIBLE_DOFD"), 4)
        self.assertGreaterEqual(identifiers.count("C-17"), 4)

    def test_v03_and_cra_state_law_gate_block_unsupported_county_and_fbpa(self):
        context = build_drafting_context(
            {"consumer_state": "FL", "consumer_address": "Miramar, Broward County, Florida", "defendants": [{"name": "Trans Union, LLC"}]},
            {"district": "Northern District of Georgia"},
        )
        issues = validate_complaint_text(
            "1. Plaintiff resides in Fulton County, Georgia.\nCOUNT I\nViolation of O.C.G.A. § 10-1-390.\n",
            context=context,
        )
        self.assertTrue(any(issue.startswith("V-03") for issue in issues))
        self.assertTrue(any(issue.startswith("V-04") for issue in issues))

    def test_v15_blocks_premature_dispute_deadline_allegation(self):
        context = build_drafting_context(
            {
                "post_dispute_disclosure_date": "07/22/2026",
                "online_disputes": [{"receipt_date": "06/26/2026"}],
            },
            {"district": "Northern District of Georgia"},
        )
        issues = validate_complaint_text(
            "1. On 06/26/2026, Defendant committed a 30-day violation by failing to reinvestigate.\n",
            context=context,
        )
        self.assertTrue(any(issue.startswith("V-15") for issue in issues))

    def test_v08_requires_structured_audit_findings_in_facts(self):
        issues = validate_complaint_text(
            "1. A generic allegation without source facts.",
            findings=[{"id": "C-01", "account": "Possible Financial", "detail": "Balance conflict."}],
        )
        self.assertTrue(any(issue.startswith("V-08") for issue in issues))

    def test_ndga_docx_has_one_caption_header_14_point_body_and_bar_number(self):
        buffer = generate_complaint_docx(
            "1. A factual paragraph.\nSIGNATURE BLOCK\n____ day of ____________, 2026",
            plaintiff_name="Labrenda Brown",
            defendant_names=["TransUnion LLC"],
        )
        doc = Document(io.BytesIO(buffer.getvalue()))
        headers = [paragraph.text for paragraph in doc.paragraphs if "IN THE UNITED STATES DISTRICT COURT" in paragraph.text]
        self.assertEqual(len(headers), 1)
        body = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("1."))
        self.assertFalse(any(run.bold or run.underline for run in body.runs if run.text.strip()))
        self.assertTrue(any("Georgia Bar No. [______]" in paragraph.text for paragraph in doc.paragraphs))
        self.assertTrue(any(run.font.size.pt == 14 for paragraph in doc.paragraphs for run in paragraph.runs if run.text.strip()))
        self.assertIn("TRANS UNION, LLC", "\n".join(paragraph.text for cell in doc.tables[0].rows[0].cells for paragraph in cell.paragraphs))


if __name__ == "__main__":
    unittest.main()
